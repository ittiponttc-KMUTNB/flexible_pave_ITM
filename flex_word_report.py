"""
flex_word_report.py — Word Report ฉบับเต็ม สำหรับ Flexible Pavement Design
AASHTO 1993 · ภาควิชาครุศาสตร์โยธา มจพ.

รูปแบบตาม flexible_consult ตัวอย่าง:
- Font: TH SarabunPSK 15pt body, 11pt equation (Times New Roman)
- Page: A4, margin L=3.2 R=3.2 T=2.5 B=2.5 cm
- Zero-Width Space (pythainlp) สำหรับ Thai line breaking
- Heading 2 = section หลัก, Heading 3 = subsection
"""

from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from pythainlp.tokenize import word_tokenize as _wt
    def _z(text: str) -> str:
        if not text: return text
        tokens = _wt(str(text), engine='newmm', keep_whitespace=True)
        return '\u200b'.join(tokens)
except ImportError:
    def _z(text: str) -> str:
        return str(text)

# ── Constants ──────────────────────────────────────────────
FN    = 'TH SarabunPSK'
FN_EQ = 'Times New Roman'
FS    = 15   # body
FS_H2 = 16   # heading 2
FS_H3 = 15   # heading 3
FS_EQ = 11   # equation
FS_TB = 15   # table
HDR_COLOR = 'BDD7EE'
PASS_RGB  = RGBColor(0x00, 0x70, 0x00)
FAIL_RGB  = RGBColor(0xC0, 0x00, 0x00)

# ============================================================
# Helpers
# ============================================================
def _tf(run, fn=FN, size=FS, bold=False, italic=False, color=None):
    run.font.name = fn
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    try:
        run._element.rPr.rFonts.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs', fn)
    except Exception:
        pass
    if color:
        run.font.color.rgb = color


def _para(doc, text='', fn=FN, size=FS, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
          space_before=4, space_after=4, indent_cm=0) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    if text:
        r = p.add_run(_z(str(text)))
        _tf(r, fn, size, bold, italic, color)
    return p


def _mixed_para(doc, runs_spec, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=6, space_after=4, indent_cm=1.25):
    """
    runs_spec: list of (text, bold) tuples
    """
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    for text, bold in runs_spec:
        r = p.add_run(_z(str(text)))
        _tf(r, FN, FS, bold)
    return p


def _eq_para(doc, text, bold=False, italic=True,
             space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    r = p.add_run(str(text))
    _tf(r, FN_EQ, FS_EQ, bold, italic)
    return p


def _hdr_shade(cell, fill=HDR_COLOR):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell(cell, text, fn=FN, size=FS_TB, bold=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, color=None, shade=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(_z(str(text)))
    _tf(r, fn, size, bold, color=color)
    if shade:
        _hdr_shade(cell, shade)


def _heading(doc, num_text, title_text, level=2):
    """Heading ที่มีเลขหัวข้อ + ชื่อ"""
    h = doc.add_heading('', level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(0)
    fs_h = FS_H2 if level == 2 else FS_H3
    r = h.add_run(f'{num_text}  {title_text}')
    _tf(r, FN, fs_h, bold=(level==2))
    return h


def _caption(doc, text):
    """คำบรรยายตาราง/รูป — bold, center"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    r = p.add_run(_z(text))
    _tf(r, FN, FS, bold=True)
    return p


def _short_mat(name: str) -> str:
    for prefix in ['ผิวทางลาดยาง ', 'พื้นทาง', 'รองพื้นทาง']:
        if name.startswith(prefix):
            return name[len(prefix):]
    return name


# ============================================================
# Main Generator
# ============================================================
def create_flex_word_report(
    project_name: str,
    designer: str,
    W18: float,
    sn_used: float,
    reliability: int,
    Zr: float,
    So: float,
    p0: float,
    pt: float,
    cbr: float,
    mr_sub: float,
    calc_results: dict,
    design_check: dict,
    fig=None,
    # Report settings
    section_num: str = '4.4',
    fig_num: str = '4-8',
    section_title: str = 'การออกแบบผิวทางลาดยาง (Flexible Pavement)',
    tbl_param_num: str = '4-8',
    tbl_mat_num: str = '4-9',
    tbl_sn_num: str = '4-10',
    tbl_param_caption: str = 'ค่าพารามิเตอร์ที่ใช้ในการออกแบบผิวทางยืดหยุ่น',
    tbl_mat_caption: str = 'ค่าสัมประสิทธิ์และค่าโมดูลัสของวัสดุโครงสร้างชั้นทาง',
    tbl_sn_caption: str = 'สรุปผลการคำนวณ Structural Number ของโครงสร้างชั้นทาง',
    fig_caption: str = 'รูปตัดโครงสร้างชั้นทางที่ออกแบบ',
    n_lanes: int = 4,
) -> BytesIO:

    doc = Document()

    # ── Page A4 ──────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3.2)
    sec.right_margin  = Cm(3.2)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    doc.styles['Normal'].font.name = FN
    doc.styles['Normal'].font.size = Pt(FS)

    delta_psi = round(p0 - pt, 1)
    sn_req    = calc_results.get('total_sn_required') or 0.0
    sn_prov   = calc_results.get('total_sn_provided', 0.0)
    passed    = design_check.get('passed', False)
    margin    = round(sn_prov - sn_req, 3)
    ratio     = round(sn_prov / sn_req, 3) if sn_req > 0 else 0.0
    layers    = calc_results.get('layers', [])
    total_d   = sum(L['design_thickness_cm'] for L in layers)
    n_layers  = len(layers)

    # sub sections
    s1 = f'{section_num}.1'
    s2 = f'{section_num}.2'
    s3 = f'{section_num}.3'
    s4 = f'{section_num}.4'
    s5 = f'{section_num}.5'

    # ──────────────────────────────────────────────────────────
    # Heading หลัก
    # ──────────────────────────────────────────────────────────
    _heading(doc, section_num, section_title, level=2)

    # ── Intro paragraph (mixed bold/normal runs) ──────────────
    result_word  = 'ผ่านเกณฑ์' if passed else 'ไม่ผ่านเกณฑ์'
    n_lanes_half = n_lanes // 2
    _mixed_para(doc, [
        (f'รูปแบบของถนนลาดยางในโครงการนี้เป็นถนน ', False),
        (f'{n_lanes}', True),
        (' ช่องจราจร ', False),
        (f'2 ทิศทาง (ไป-กลับ)', True),
        (' การออกแบบโครงสร้างถนนแบบยืดหยุ่น (Flexible Pavement) ใช้วิธี AASHTO 1993 '
         'Guide for Design of Pavement Structures โดยพิจารณาจำนวนเพลาสะสม ESALs '
         'ความน่าเชื่อถือ และผลสัมปสิทธิ์ชั้นทาง สำหรับโครงการนี้กำหนดให้ค่า W18 = ', False),
        (f'{W18:,.0f}', True),
        (' 18-kip ESALs ที่ระดับความน่าเชื่อถือ (Reliability) = ', False),
        (f'{reliability}', True),
        (' % โดยมีดินเดิมค่า CBR = ', False),
        (f'{cbr:.1f}', True),
        (f' % (Mr = ', False),
        (f'{mr_sub:,.0f}', True),
        (f' psi) ผลการออกแบบได้โครงสร้างชั้นทาง ', False),
        (f'{n_layers}', True),
        (' ชั้น ที่ SN_required = ', False),
        (f'{sn_req:.2f}', True),
        (' และ SN_provided = ', False),
        (f'{sn_prov:.2f}', True),
        (' ความหนารวม ', False),
        (f'{total_d:.0f}', True),
        (f' ซม. การออกแบบ', False),
        (result_word, True),
        (f' ดังแสดงผลการวิเคราะห์ในตารางที่ ', False),
        (tbl_param_num, True),
        (' และตารางที่ ', False),
        (tbl_mat_num, True),
        (' และรูปที่ ', False),
        (fig_num, True),
    ], space_before=6, space_after=4, indent_cm=1.25)

    # ══════════════════════════════════════════════════════════
    # 4.4.1 วิธีการออกแบบ
    # ══════════════════════════════════════════════════════════
    _heading(doc, s1, 'วิธีการออกแบบ', level=3)
    _para(doc,
        'การออกแบบโครงสร้างถนนใช้วิธี AASHTO 1993 Guide for Design of Pavement Structures '
        'ตามมาตรฐานกรมทางหลวง โดยใช้สมการหลักดังนี้',
        space_before=4, space_after=4, indent_cm=1.25)
    _eq_para(doc,
        'log10(W18) = Zr\u00b7So + 9.36\u00b7log10(SN+1) - 0.20\n'
        '                   + log10(\u0394PSI/2.7) / [0.4 + 1094/(SN+1)^5.19] + 2.32\u00b7log10(Mr) - 8.07',
        space_before=0, space_after=4)

    # ══════════════════════════════════════════════════════════
    # 4.4.2 ข้อมูลนำเข้า
    # ══════════════════════════════════════════════════════════
    _heading(doc, s2, 'ข้อมูลนำเข้า (Design Inputs)', level=3)
    _para(doc,
        'ในการออกแบบโครงสร้างถนนยืดหยุ่น การกำหนดค่าพารามิเตอร์นำเข้า (Design Inputs) '
        f'ถือเป็นขั้นตอนสำคัญ ซึ่งในโครงการนี้ได้กำหนดค่าพารามิเตอร์ต่างๆ '
        f'ดังแสดงในตารางที่ {tbl_param_num}',
        space_before=4, space_after=4, indent_cm=1.25)
    _caption(doc, f'ตารางที่ {tbl_param_num}  {tbl_param_caption}')

    param_rows = [
        ('Design ESALs (W\u2081\u2088)',       f'{W18:,.0f}',     '18-kip ESAL'),
        (f'Structural Number (SN = {sn_used:.1f})', f'{sn_used:.1f}', '\u2014'),
        ('Reliability (R)',                    f'{reliability}',   '%'),
        ('Z\u1d63',                            f'{Zr:.3f}',        '\u2014'),
        ('S\u2080',                            f'{So:.2f}',        '\u2014'),
        ('P\u2080',                            f'{p0:.1f}',        '\u2014'),
        ('P\u209c',                            f'{pt:.1f}',        '\u2014'),
        ('\u0394PSI',                          f'{delta_psi:.1f}', '\u2014'),
        ('CBR ดินเดิม',                       f'{cbr:.1f}',       '%'),
        ('M\u1d63',                            f'{mr_sub:,.0f}',   'psi'),
    ]
    cw0 = [5.08*Cm(1), 5.08*Cm(1), 5.08*Cm(1)]
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['พารามิเตอร์', 'ค่า', 'หน่วย']):
        _cell(t1.rows[0].cells[j], h, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, shade=HDR_COLOR)
    for pm, val, unit in param_rows:
        row = t1.add_row()
        _cell(row.cells[0], pm)
        _cell(row.cells[1], val,  align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # 4.4.3 คุณสมบัติวัสดุ
    # ══════════════════════════════════════════════════════════
    _heading(doc, s3, 'คุณสมบัติวัสดุชั้นทาง', level=3)
    _para(doc,
        'วัสดุโครงสร้างชั้นทางแต่ละชนิดมีค่าสัมประสิทธิ์ชั้นทาง (Layer Coefficient) '
        'และค่าโมดูลัสความยืดหยุ่น (Resilient Modulus) ดังแสดงในตารางที่ '
        f'{tbl_mat_num}',
        space_before=4, space_after=4, indent_cm=1.25)
    _caption(doc, f'ตารางที่ {tbl_mat_num}  {tbl_mat_caption}')

    t2 = doc.add_table(rows=1, cols=6)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['ชั้น','วัสดุ','a\u1d62','m\u1d62','M\u1d63 (psi)','E (MPa)']):
        _cell(t2.rows[0].cells[j], h, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, shade=HDR_COLOR)
    for L in layers:
        row = t2.add_row()
        _cell(row.cells[0], str(L['layer_no']),  align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[1], _short_mat(L['material']))
        _cell(row.cells[2], f'{L["a_i"]:.2f}',  align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[3], f'{L["m_i"]:.2f}',  align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[4], f'{L["mr_psi"]:,}', align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell(row.cells[5], f'{L["mr_mpa"]:,}', align=WD_ALIGN_PARAGRAPH.RIGHT)

    # AC sublayer
    for L in layers:
        ac_sub = L.get('ac_sublayers') or L.get('ac_sub')
        if ac_sub and isinstance(ac_sub, dict):
            doc.add_paragraph()
            _para(doc, 'รายละเอียดชั้นย่อยผิวทาง AC:', bold=True,
                  space_before=4, space_after=2)
            t_sub = doc.add_table(rows=1, cols=3)
            t_sub.style = 'Table Grid'
            for j, h in enumerate(['ชั้นย่อย','ความหนา (cm)','ความหนา (mm)']):
                _cell(t_sub.rows[0].cells[j], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, shade=HDR_COLOR)
            for label, key in [
                ('ผิวทาง Wearing Course', 'wearing_cm'),
                ('รองผิวทาง Binder Course', 'binder_cm'),
                ('พื้นทาง Base Course', 'base_cm'),
            ]:
                t_cm = ac_sub.get(key, 0)
                row  = t_sub.add_row()
                _cell(row.cells[0], label)
                _cell(row.cells[1], f'{t_cm:.1f}', align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell(row.cells[2], f'{t_cm*10:.0f}', align=WD_ALIGN_PARAGRAPH.CENTER)
            break
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # 4.4.4 ขั้นตอนการคำนวณความหนาชั้นทาง
    # ══════════════════════════════════════════════════════════
    _heading(doc, s4, 'ขั้นตอนการคำนวณความหนาชั้นทาง', level=3)
    _para(doc,
        'การคำนวณความหนาขั้นต่ำของแต่ละชั้น ใช้หลักการว่า Structural Number (SN) '
        'ที่จุดใดๆ ของโครงสร้างต้องมากกว่าหรือเท่ากับค่า SN ที่คำนวณได้จากสมการ AASHTO 1993',
        space_before=4, space_after=4, indent_cm=1.25)

    for L in layers:
        ln      = L['layer_no']
        a_i     = L['a_i']
        m_i     = L['m_i']
        d_in    = L['design_thickness_inch']
        d_cm    = L['design_thickness_cm']
        sn_at   = L['sn_required_at_layer']
        d_min   = L['min_thickness_cm']
        sn_cont = L['sn_contribution']
        sn_cum  = L['cumulative_sn']
        is_ok   = L['is_ok']

        doc.add_paragraph()
        # ชื่อชั้น
        _para(doc, f'ชั้นที่ {ln}: {_short_mat(L["material"])}',
              bold=True, space_before=6, space_after=4)
        # ข้อมูลวัสดุ
        _para(doc, 'ข้อมูลวัสดุ:', bold=True, space_before=4, space_after=2)
        _para(doc,
            f'\u2022 Mr = {L["mr_psi"]:,} psi  =  {L["mr_mpa"]:,} MPa\n'
            f'\u2022 Layer Coefficient (a{ln}) = {a_i:.2f}\n'
            f'\u2022 Drainage Coefficient (m{ln}) = {m_i:.2f}',
            space_before=2, space_after=4)

        # การคำนวณ SN
        _para(doc, 'การคำนวณ SN:', bold=True, space_before=4, space_after=2)
        _eq_para(doc, f'จากสมการ AASHTO 1993:   SN_{ln} = {sn_at:.2f}',
                 bold=True, space_before=0, space_after=4)

        # ความหนาขั้นต่ำ
        _para(doc, 'การคำนวณความหนาขั้นต่ำ:', bold=True, space_before=4, space_after=2)
        if ln == 1:
            _eq_para(doc, f'D_{ln} >= SN_{ln} / (a_{ln} \u00d7 m_{ln})',
                     bold=False, space_before=0, space_after=4)
        else:
            _eq_para(doc,
                f'D_{ln} >= (SN_{ln} \u2212 SN_{ln-1}) / (a_{ln} \u00d7 m_{ln})',
                bold=False, space_before=0, space_after=4)

        # เลือกใช้ความหนา
        _para(doc, 'เลือกใช้ความหนา:', bold=True, space_before=4, space_after=2)
        _eq_para(doc, f'D_{ln}(design)  =  {d_cm:.0f} cm  ({d_in:.2f} in)',
                 bold=True, italic=False, space_before=0, space_after=4)

        # SN contribution
        _para(doc, 'SN contribution:', bold=True, space_before=4, space_after=2)
        _eq_para(doc,
            f'\u0394SN_{ln} = a_{ln} \u00d7 D_{ln} \u00d7 m_{ln}'
            f'  =  {a_i:.2f} \u00d7 {d_in:.2f} \u00d7 {m_i:.2f}  =  {sn_cont:.3f}',
            bold=False, space_before=0, space_after=2)
        _eq_para(doc, f'\u03a3SN  =  {sn_cum:.2f}',
                 bold=True, italic=False, space_before=0, space_after=4)

        # สถานะ
        if is_ok:
            status_txt = f'\u2713 OK  \u2014  ความหนาเพียงพอ'
            col = PASS_RGB
        else:
            short_need = d_min - d_cm
            status_txt = f'\u2717 NG  \u2014  ต้องเพิ่มความหนาอีก {short_need:.1f} cm'
            col = FAIL_RGB
        p_st = _para(doc, f'สถานะ:  {status_txt}', bold=True,
                     space_before=4, space_after=4)
        for r in p_st.runs:
            r.font.color.rgb = col

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # 4.4.5 สรุปการคำนวณ SN
    # ══════════════════════════════════════════════════════════
    _heading(doc, s5, 'สรุปการคำนวณ Structural Number', level=3)
    _para(doc,
        f'เมื่อได้ทำการคำนวณความหนาของแต่ละชั้นทางตามขั้นตอนที่กล่าวมาข้างต้นแล้ว '
        f'สามารถสรุปผลการคำนวณ Structural Number ได้ดังตารางที่ {tbl_sn_num}',
        space_before=4, space_after=4, indent_cm=1.25)
    _caption(doc, f'ตารางที่ {tbl_sn_num}  {tbl_sn_caption}')

    t3 = doc.add_table(rows=1, cols=8)
    t3.style = 'Table Grid'
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['ชั้น','วัสดุ','a\u1d62','m\u1d62',
                            'D\u1d62 (นิ้ว)','D\u1d62 (ซม.)','\u0394SN\u1d62','\u03a3SN']):
        _cell(t3.rows[0].cells[j], h, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, shade=HDR_COLOR)
    for L in layers:
        row = t3.add_row()
        _cell(row.cells[0], str(L['layer_no']),                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[1], _short_mat(L['material']))
        _cell(row.cells[2], f'{L["a_i"]:.2f}',                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[3], f'{L["m_i"]:.2f}',                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[4], f'{L["design_thickness_inch"]:.2f}', align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[5], f'{L["design_thickness_cm"]:.0f}',   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[6], f'{L["sn_contribution"]:.3f}',       align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[7], f'{L["cumulative_sn"]:.2f}',         align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ── ผลการตรวจสอบ ──────────────────────────────────────────
    _para(doc, 'ผลการตรวจสอบการออกแบบ', bold=True,
          space_before=4, space_after=4)

    result_rows = [
        ('SN Required (จากสมการ AASHTO)', f'{sn_req:.2f}'),
        ('SN Provided (จากชั้นทาง)',       f'{sn_prov:.2f}'),
        ('Safety Margin (SN_provided \u2212 SN_required)', f'{margin:+.3f}'),
        ('Ratio (SN_provided / SN_required)', f'{ratio:.3f}'),
        ('ผลการตรวจสอบ', 'ผ่าน (OK)' if passed else 'ไม่ผ่าน (NG)'),
    ]
    t4 = doc.add_table(rows=len(result_rows), cols=2)
    t4.style = 'Table Grid'
    t4.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (param, val) in enumerate(result_rows):
        is_last  = (i == len(result_rows) - 1)
        col      = (PASS_RGB if passed else FAIL_RGB) if is_last else None
        _cell(t4.rows[i].cells[0], param, bold=is_last)
        _cell(t4.rows[i].cells[1], val,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=is_last, color=col)
        if is_last:
            sh = 'C6EFCE' if passed else 'FFC7CE'
            _hdr_shade(t4.rows[i].cells[0], sh)
            _hdr_shade(t4.rows[i].cells[1], sh)
    doc.add_paragraph()

    # สรุป
    summary = (
        f'สรุป: การออกแบบผ่านเกณฑ์ เนื่องจาก SN_provided ({sn_prov:.2f}) '
        f'\u2265 SN_required ({sn_req:.2f})'
        if passed else
        f'สรุป: การออกแบบไม่ผ่านเกณฑ์ เนื่องจาก SN_provided ({sn_prov:.2f}) '
        f'< SN_required ({sn_req:.2f}) กรุณาปรับเพิ่มความหนาชั้นทาง'
    )
    _para(doc, summary, bold=True,
          color=PASS_RGB if passed else FAIL_RGB,
          space_before=4, space_after=4)
    doc.add_paragraph()

    # ── ตารางสรุปโครงสร้างชั้นทาง + รูปตัดขวาง (merge ในเซลล์แรก) ──
    _caption(doc, f'รูปที่ {fig_num}  {fig_caption}')
    _caption(doc, f'ตารางที่ {tbl_sn_num}  {tbl_sn_caption}')

    # เตรียม png bytes ของรูป (ถ้ามี)
    _fig_buf = None
    if fig is not None:
        try:
            import matplotlib.pyplot as plt
            _b = BytesIO()
            fig.savefig(_b, format='png', dpi=150,
                        bbox_inches='tight', facecolor='white')
            _b.seek(0)
            _fig_buf = _b
            plt.close(fig)
        except Exception:
            _fig_buf = None

    # สร้างตาราง: header row + n layer rows + 1 subgrade row
    n_data_rows = len(layers) + 1          # layers + subgrade
    t5 = doc.add_table(rows=1 + n_data_rows, cols=3)
    t5.style = 'Table Grid'
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── Header row ──
    for j, h in enumerate(['รายละเอียด', 'หนา (ซม.)', 'ชนิดวัสดุ']):
        _cell(t5.rows[0].cells[j], h, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, shade=HDR_COLOR)

    # ── Data rows (layers) ──
    for idx, L in enumerate(layers):
        row = t5.rows[1 + idx]
        _cell(row.cells[0], '')           # จะ vMerge ทีหลัง
        _cell(row.cells[1], f'{L["design_thickness_cm"]:.0f}',
              align=WD_ALIGN_PARAGRAPH.CENTER)
        from flex_engine import MATERIALS as _MAT
        _eng = _MAT.get(L['material'], {}).get('english_name', _short_mat(L['material']))
        _cell(row.cells[2], _eng)

    # ── Subgrade row ──
    sub_row = t5.rows[1 + len(layers)]
    _cell(sub_row.cells[0], '')           # จะ vMerge ทีหลัง
    _cell(sub_row.cells[1], 'Existing', align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(sub_row.cells[2],
          f'Earth Embankment / Subgrade\nCBR\u2265{cbr:.1f} %')

    # ── vMerge เซลล์แรกของแถวข้อมูลทั้งหมด (แถว 1 ถึง n_data_rows) ──
    def _vmerge_start(cell):
        """กำหนดให้ cell เป็นจุดเริ่ม vMerge"""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vm   = OxmlElement('w:vMerge')
        vm.set(qn('w:val'), 'restart')
        tcPr.append(vm)

    def _vmerge_cont(cell):
        """กำหนดให้ cell ต่อเนื่องจาก vMerge"""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vm   = OxmlElement('w:vMerge')
        # ไม่ set val = ต่อเนื่อง
        tcPr.append(vm)

    _vmerge_start(t5.rows[1].cells[0])
    for r_idx in range(2, 1 + n_data_rows):
        _vmerge_cont(t5.rows[r_idx].cells[0])

    # ── แทรกรูปในเซลล์แรกของแถวที่ 1 (vMerge restart) ──
    first_cell = t5.rows[1].cells[0]
    first_cell.text = ''
    p_img = first_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _fig_buf is not None:
        try:
            r_img = p_img.add_run()
            r_img.add_picture(_fig_buf, width=Inches(2.2))
        except Exception as e:
            r_img = p_img.add_run(f'[รูปไม่แสดง: {e}]')
            _tf(r_img, FN, FS - 2)

    doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────
    _para(doc,
        f'พัฒนาโดย รศ.ดร.อิทธิพล มีผล \u00b7 ภาควิชาครุศาสตร์โยธา \u00b7 มจพ. \u00b7 '
        'Flexible Pavement Design V1 \u00b7 AASHTO 1993',
        size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=0)

    buf_out = BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    return buf_out
